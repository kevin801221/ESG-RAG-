# tests/test_rag.py

import os
import sys
import pytest
from pathlib import Path
import pandas as pd
from dotenv import load_dotenv
import docx
from src.extractors.metric_extractor import ESGMetricExtractor
from src.extractors.text_extractor import DocumentProcessor
from src.db.vector_store import KDBAIStore
from src.config.config_loader import load_config, load_metrics_schema

pytestmark = pytest.mark.asyncio

class TestRAGSystem:
    @pytest.fixture(scope="class")
    def setup(self):
        """設置測試環境"""
        # 載入環境變數和配置
        load_dotenv()
        config = load_config()
        metrics_schema = load_metrics_schema()
        
        # 初始化組件
        extractor = ESGMetricExtractor(config)
        doc_processor = DocumentProcessor(config)
        
        return {
            'config': config,
            'metrics_schema': metrics_schema,
            'extractor': extractor,
            'doc_processor': doc_processor
        }
    
    def create_test_docx(self):
        """創建測試用的 Word 文檔"""
        doc = docx.Document()
        
        # 添加標題
        doc.add_heading('永續報告書 2023', 0)
        
        # 添加環境指標段落
        doc.add_heading('環境指標', 1)
        doc.add_paragraph('溫室氣體排放量：12,345 公噸CO2e')
        doc.add_paragraph('用水量：45,678 立方米')
        
        # 添加社會指標段落
        doc.add_heading('社會指標', 1)
        doc.add_paragraph('員工培訓時數：42.5 小時/人')
        doc.add_paragraph('整體薪酬水準：85.6%')
        
        # 添加表格
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # 填充表格
        cells = table.rows[0].cells
        cells[0].text = '指標'
        cells[1].text = '數值'
        
        cells = table.rows[1].cells
        cells[0].text = '溫室氣體排放量'
        cells[1].text = '12,345 公噸CO2e'
        
        cells = table.rows[2].cells
        cells[0].text = '員工培訓時數'
        cells[1].text = '42.5 小時/人'
        
        # 保存文檔
        test_file = Path("test_doc.docx")
        doc.save(str(test_file))
        return test_file

    @pytest.mark.asyncio
    async def test_document_processing(self, setup):
        """測試文檔處理功能"""
        # 創建測試文檔
        test_file = self.create_test_docx()
        
        try:
            # 處理文檔
            chunks, metadata = setup['doc_processor'].process_document(test_file)
            
            # 驗證結果
            assert len(chunks) > 0, "應該產生至少一個文本塊"
            assert metadata['file_name'] == "test_doc.docx"
            assert metadata['file_type'] == '.docx'
            assert 'process_time' in metadata
            assert metadata['tables_found'] > 0, "應該找到至少一個表格"
            
            print("\n✅ 文檔處理測試通過")
            print(f"生成了 {len(chunks)} 個文本塊")
            print(f"找到 {metadata['tables_found']} 個表格")
            
            return chunks
            
        finally:
            # 清理測試文件
            if test_file.exists():
                test_file.unlink()

    @pytest.mark.asyncio
    async def test_metric_extraction(self, setup):
        """測試指標提取功能"""
        # 使用上一個測試的文檔塊
        chunks = await self.test_document_processing(setup)
        
        # 提取指標
        results_df = await setup['extractor'].process_document(
            chunks=chunks,
            document_id="test_doc",
            metadata={'source': 'test'}
        )
        
        # 驗證結果
        assert isinstance(results_df, pd.DataFrame)
        assert len(results_df) > 0
        assert '溫室氣體排放量' in results_df['項目'].values
        assert '用水量' in results_df['項目'].values
        
        print("\n✅ 指標提取測試通過")
        print("\n提取到的指標：")
        print(results_df)

    @pytest.mark.asyncio
    async def test_retrieval(self, setup):
        """測試檢索功能"""
        # 準備測試查詢
        test_queries = [
            "公司的溫室氣體排放量是多少？",
            "員工平均培訓時數？",
            "用水量數據"
        ]
        
        # 對每個查詢進行測試
        for query in test_queries:
            relevant_chunks = await setup['extractor'].vector_store.similarity_search(
                query,
                k=2
            )
            
            # 驗證結果
            assert len(relevant_chunks) > 0, f"查詢'{query}'應該返回相關結果"
            
            print(f"\n查詢: {query}")
            print(f"找到 {len(relevant_chunks)} 個相關文本塊")
            for i, chunk in enumerate(relevant_chunks, 1):
                print(f"\n相關文本 {i}:")
                print(chunk['text'])
                if 'score' in chunk:
                    print(f"相似度分數: {chunk['score']:.3f}")
                
        print("\n✅ 檢索測試通過")

def main():
    """執行所有測試"""
    print("\n=== 開始測試 RAG 系統 ===\n")
    
    # 設置 pytest 參數
    pytest_args = [
        '-v',
        '--asyncio-mode=auto',
        'tests/test_rag.py'
    ]
    
    # 運行測試
    pytest.main(pytest_args)

if __name__ == "__main__":
    main()