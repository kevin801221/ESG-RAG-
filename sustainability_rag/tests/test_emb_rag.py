# tests/test_emb_rag.py

import os
import pytest
import pandas as pd
import numpy as np
from pathlib import Path
import docx
from dotenv import load_dotenv
import nest_asyncio

# 在最開始就應用 nest_asyncio
nest_asyncio.apply()

# 在最開始就載入環境變數
load_dotenv(Path(__file__).parent.parent / '.env')

from llama_index.embeddings.openai import OpenAIEmbedding
from src.extractors.text_extractor import DocumentProcessor
from src.extractors.metric_extractor import ESGMetricExtractor
from src.config.config_loader import load_config

@pytest.fixture(scope="session", autouse=True)
def setup_environment():
    """確保環境變數正確設置"""
    # 驗證必要的環境變數
    required_vars = [
        'LLAMA_CLOUD_API_KEY',
        'OPENAI_API_KEY',
        'KDBAI_ENDPOINT',
        'KDBAI_API_KEY'
    ]
    
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    if missing_vars:
        pytest.fail(f"Missing required environment variables: {missing_vars}")
        
    print("\n=== Environment Setup ===")
    for var in required_vars:
        value = os.getenv(var)
        print(f"{var}: {'[SET]' if value else '[MISSING]'}")

@pytest.fixture
def test_doc_path(tmp_path):
    """創建測試用的 Word 文檔"""
    doc_path = tmp_path / "test_esg_report.docx"
    doc = docx.Document()
    
    # 添加測試內容
    doc.add_heading('ESG 報告測試文件', 0)
    
    doc.add_heading('環境指標', 1)
    doc.add_paragraph('2023年溫室氣體排放量為12,345公噸CO2e，相較2022年減少5%。')
    
    doc.add_heading('社會指標', 1)
    doc.add_paragraph('員工人數達到500人，女性主管比例為35%。')
    doc.add_paragraph('員工平均training時數為42小時/年。')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=2)
    cells = table.rows[0].cells
    cells[0].text = '指標'
    cells[1].text = '數值'
    
    cells = table.rows[1].cells
    cells[0].text = '能源使用量'
    cells[1].text = '1,234,567 kWh'
    
    cells = table.rows[2].cells
    cells[0].text = '工傷率'
    cells[1].text = '0.5%'
    
    doc.save(doc_path)
    return doc_path

@pytest.fixture
def config():
    """載入測試配置"""
    config = load_config()
    # 確保加載了 API key
    config['llamaparse'] = {
        'api_key': os.getenv('LLAMA_CLOUD_API_KEY')
    }
    return config

@pytest.mark.asyncio
async def test_document_embedding(test_doc_path, config):
    """測試文檔嵌入功能"""
    try:
        # 初始化處理器
        doc_processor = DocumentProcessor(config)
        embed_model = OpenAIEmbedding()
        
        # 處理文檔
        chunks, metadata = await doc_processor.aprocess_document(test_doc_path)
        
        print("\n=== 文檔處理結果 ===")
        print(f"生成的文本塊數: {len(chunks)}")
        print(f"找到的表格數: {metadata['tables_found']}")
        
        # 測試嵌入生成
        embeddings = await embed_model.aget_text_embedding_batch(chunks)
        
        # 驗證嵌入向量
        assert len(embeddings) == len(chunks)
        assert all(len(emb) == 1536 for emb in embeddings)
        
        print("\n=== 嵌入向量資訊 ===")
        print(f"向量維度: {len(embeddings[0])}")
        print(f"向量數量: {len(embeddings)}")
        
        return chunks, embeddings
    except Exception as e:
        pytest.fail(f"測試失敗: {str(e)}")

@pytest.mark.asyncio
async def test_rag_pipeline(test_doc_path, config):
    """測試完整的 RAG 流程"""
    try:
        # 初始化組件
        doc_processor = DocumentProcessor(config)
        metric_extractor = ESGMetricExtractor(config)
        
        # 處理文檔
        chunks, metadata = await doc_processor.aprocess_document(test_doc_path)
        
        # 提取指標
        results_df = await metric_extractor.process_document(
            chunks=chunks,
            document_id="test_doc",
            metadata={'source': 'test'}
        )
        
        print("\n=== RAG 處理結果 ===")
        print("\n提取到的指標：")
        print(results_df)
        
        # 測試相似度搜索
        test_queries = [
            "公司的溫室氣體排放量是多少？",
            "員工培訓情況如何？",
            "能源使用情況是？"
        ]
        
        print("\n=== 相似度搜索測試 ===")
        for query in test_queries:
            results = await metric_extractor.vector_store.similarity_search(
                query=query,
                k=2
            )
            print(f"\n查詢: {query}")
            print(f"找到 {len(results)} 個相關結果")
            for i, result in enumerate(results, 1):
                print(f"\n相關文本 {i}:")
                print(f"文本: {result['text'][:200]}...")
                if 'score' in result:
                    print(f"相似度分數: {result['score']:.3f}")
                    
        # 驗證結果
        assert not results_df.empty, "應該找到至少一個指標"
        assert '溫室氣體排放量' in results_df['項目'].values
        assert '員工培訓時數' in results_df['項目'].values
        
    except Exception as e:
        pytest.fail(f"測試失敗: {str(e)}")
    finally:
        # 清理資源
        metric_extractor.cleanup()

if __name__ == "__main__":
    pytest.main(['-v', __file__])